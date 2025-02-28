import os
import io
import xmlrpc.client
import streamlit as st
from dotenv import load_dotenv
from docx import Document
import datetime
from typing import Tuple, Optional, List, Dict, Any

# For removing empty paragraphs and setting footer font size
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Load environment variables
load_dotenv()

# === Odoo Credentials from Environment Variables ===
ODOO_URL = os.getenv("ODOO_URL", "https://prezlab-staging-17999869.dev.odoo.com")
ODOO_DB = os.getenv("ODOO_DB", "prezlab-staging-17999869")
ODOO_USERNAME = os.getenv("ODOO_USERNAME", "omar.elhasan@prezlab.com")
ODOO_PASSWORD = os.getenv("ODOO_PASSWORD", "1234")

# --- Odoo Connection using caching for robustness and speed ---
@st.cache_resource(show_spinner=False)
def get_odoo_connection() -> Tuple[Optional[int], Optional[xmlrpc.client.ServerProxy]]:
    try:
        common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
        uid = common.authenticate(ODOO_DB, ODOO_USERNAME, ODOO_PASSWORD, {})
        if not uid:
            st.error("Failed to authenticate with Odoo. Check credentials and database name.")
            return None, None
        models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
        return uid, models
    except Exception as e:
        st.error(f"Error connecting to Odoo: {e}")
        return None, None

def get_employee_fields(models: xmlrpc.client.ServerProxy, uid: int) -> List[str]:
    try:
        fields = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'fields_get', [], {'attributes': ['type']})
        return list(fields.keys())
    except Exception as e:
        st.error(f"Error retrieving employee fields: {e}")
        return []

def get_arabic_name(employee: Dict[str, Any]) -> str:
    # Check for the custom Arabic name field first.
    name = employee.get("x_studio_employee_arabic_name", "").strip()
    if name:
        return name
    return employee.get("name", "").strip()

def get_partner_address(models: xmlrpc.client.ServerProxy, uid: int, partner_id: int) -> str:
    try:
        fields = ["street", "street2", "city", "zip", "country_id"]
        result = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, "res.partner", "read", [[partner_id]], {"fields": fields})
        if result:
            partner = result[0]
            country = ""
            if partner.get("country_id"):
                if isinstance(partner["country_id"], (list, tuple)) and len(partner["country_id"]) > 1:
                    country = partner["country_id"][1]
                else:
                    country = str(partner["country_id"])
            address_parts = [partner.get("street", ""), partner.get("street2", ""), partner.get("city", ""), partner.get("zip", ""), country]
            address_str = ", ".join([part for part in address_parts if part])
            return address_str
        else:
            return ""
    except Exception as e:
        st.error(f"Error retrieving partner address: {e}")
        return ""

def get_arabic_partner_address(models: xmlrpc.client.ServerProxy, uid: int, partner_id: int) -> str:
    try:
        fields = ["x_studio_arabic_address"]
        result = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, "res.partner", "read", [[partner_id]], {"fields": fields})
        if result:
            partner = result[0]
            arabic_address = partner.get("x_studio_arabic_address", "")
            return arabic_address.strip() if arabic_address else ""
        else:
            return ""
    except Exception as e:
        st.error(f"Error retrieving arabic partner address: {e}")
        return ""

def get_company_registrar(models: xmlrpc.client.ServerProxy, uid: int, company_id: int) -> str:
    try:
        result = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, "res.company", "read", [[company_id]], {"fields": ["company_registry"]})
        if result:
            company = result[0]
            return company.get("company_registry", "")
        else:
            return ""
    except Exception as e:
        st.error(f"Error retrieving company registrar: {e}")
        return ""

def get_company_arabic_name(models: xmlrpc.client.ServerProxy, uid: int, company_id: int) -> str:
    try:
        result = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, "res.company", "read", [[company_id]], {"fields": ["arabic_name"]})
        if result:
            company = result[0]
            return company.get("arabic_name", "")
        else:
            return ""
    except Exception as e:
        st.error(f"Error retrieving company Arabic name: {e}")
        return ""

def get_head_people_and_culture(models: xmlrpc.client.ServerProxy, uid: int, company_id: int) -> str:
    try:
        domain = [('company_id', '=', company_id), ('job_id.name', 'ilike', 'head of people and culture')]
        head_ids = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'search', [domain])
        if head_ids:
            head_data = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'read', [head_ids[0]], {'fields': ['name']})
            if head_data:
                return head_data[0].get('name', '')
        return ""
    except Exception as e:
        st.error(f"Error retrieving head of people and culture: {e}")
        return ""

def get_head_people_and_culture_arabic(models: xmlrpc.client.ServerProxy, uid: int, company_id: int) -> str:
    try:
        domain = [('company_id', '=', company_id), ('job_id.name', 'ilike', 'head of people and culture')]
        head_ids = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'search', [domain])
        if head_ids:
            head_data = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'read', [head_ids[0]], {'fields': ['x_studio_employee_arabic_name', 'name']})
            if head_data:
                return get_arabic_name(head_data[0])
        return ""
    except Exception as e:
        st.error(f"Error retrieving head of people and culture (Arabic): {e}")
        return ""

def derive_country_from_address(address: str) -> str:
    if not address:
        return ""
    if "\n" in address:
        lines = [line.strip() for line in address.split("\n") if line.strip()]
        if lines:
            return lines[-1]
    parts = [part.strip() for part in address.split(",") if part.strip()]
    if parts:
        return parts[-1]
    return ""

def get_employee_by_id(models: xmlrpc.client.ServerProxy, uid: int, identification_no: str) -> Optional[Dict[str, Any]]:
    try:
        identification_no = identification_no.strip()
        available_fields = get_employee_fields(models, uid)
        if "identification_id" in available_fields:
            search_domain = [('identification_id', '=', identification_no)]
        else:
            st.error("Field 'identification_id' does not exist in Odoo.")
            return None
        # Include the custom joining date, contract end date, and department fields in the fields to read
        fields_to_read = [
            'id', 'name', 'job_title', 'x_studio_joining_date', 'x_studio_employee_arabic_name',
            'identification_id', 'company_id', 'address_id', 'x_studio_contract_end_date', 'department_id'
        ]
        employee_ids = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'search', [search_domain])
        if not employee_ids:
            st.warning("No employee found with the provided identification number.")
            return None
        employee_data_list = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.employee', 'read',
                                                 [employee_ids], {'fields': fields_to_read})
        if not employee_data_list:
            st.warning("Employee data retrieval failed.")
            return None
        if len(employee_data_list) > 1:
            names = [emp.get("name", "") for emp in employee_data_list]
            st.info("Multiple employees found with the same identification number: " + ", ".join(names) + ". Using the first match.")
        employee = employee_data_list[0]
        try:
            contracts = models.execute_kw(ODOO_DB, uid, ODOO_PASSWORD, 'hr.contract', 'search_read',
                                            [[('employee_id', '=', employee['id'])]],
                                            {'fields': ['wage'], 'limit': 1})
        except xmlrpc.client.Fault:
            st.warning("User does not have access to hr.contract records. Default wage will be used.")
            contracts = []
        wage = contracts[0].get('wage', 0.0) if contracts else 0.0

        # Process the joining date from the custom field x_studio_joining_date
        joining_date_raw = employee.get('x_studio_joining_date', '')
        joining_date_str = ""
        if joining_date_raw:
            try:
                joining_date_dt = datetime.datetime.strptime(joining_date_raw, "%Y-%m-%d")
                joining_date_str = joining_date_dt.strftime("%d/%m/%Y")
            except Exception:
                joining_date_str = joining_date_raw

        # Process the contract end date from the custom field
        contract_end_date_raw = employee.get('x_studio_contract_end_date', '')
        contract_end_date_str = ""
        if contract_end_date_raw:
            try:
                contract_end_date_dt = datetime.datetime.strptime(contract_end_date_raw, "%Y-%m-%d")
                contract_end_date_str = contract_end_date_dt.strftime("%d/%m/%Y")
            except Exception:
                contract_end_date_str = contract_end_date_raw

        # Retrieve the department from the department_id field
        department_field = employee.get("department_id")
        department = ""
        if department_field and isinstance(department_field, (list, tuple)):
            department = department_field[1] if len(department_field) > 1 else str(department_field[0])
        elif department_field:
            department = str(department_field)

        arabic_name = get_arabic_name(employee)
        company = ""
        company_field = employee.get("company_id")
        company_registrar = ""
        company_arabic_name = ""
        head_people_culture = ""
        head_people_culture_arabic = ""
        if company_field and isinstance(company_field, (list, tuple)) and len(company_field) > 0:
            company = company_field[1] if len(company_field) > 1 else str(company_field[0])
            company_registrar = get_company_registrar(models, uid, company_field[0])
            company_arabic_name = get_company_arabic_name(models, uid, company_field[0])
            if not company_arabic_name:
                company_arabic_name = company
            head_people_culture = get_head_people_and_culture(models, uid, company_field[0])
            head_people_culture_arabic = get_head_people_and_culture_arabic(models, uid, company_field[0])
        elif company_field:
            company = str(company_field)
        
        # Default values if Head of People & Culture is not found
        if not head_people_culture:
            head_people_culture = "Faisal Abdullah AlMamun"
        if not head_people_culture_arabic:
            head_people_culture_arabic = "فيصل عبدالله المأمون"
        
        # Retrieve the work address from the address_id field.
        # If the field is a Many2one (list/tuple), retrieve the partner record;
        # otherwise, assume the field contains the work address as a string.
        work_address = ""
        arabic_work_address = ""
        address_field = employee.get("address_id")
        if address_field:
            if isinstance(address_field, (list, tuple)):
                if len(address_field) > 0:
                    partner_id = address_field[0]
                    work_address = get_partner_address(models, uid, partner_id)
                    arabic_work_address = get_arabic_partner_address(models, uid, partner_id)
            else:
                work_address = str(address_field)
                arabic_work_address = ""
                
        company_country = derive_country_from_address(work_address)
        return {
            'id': employee.get('id', ''),
            'name': employee.get('name', '').strip(),
            'first_name': employee.get('name', '').split()[0] if employee.get('name') else '',
            'job_title': employee.get('job_title', '').strip(),
            'identification': employee.get('identification_id', '').strip(),
            'wage': wage,
            'joining_date': joining_date_str,
            'contract_end_date': contract_end_date_str,
            'department': department,
            'arabic_name': arabic_name,
            'company': company,
            'work_address': work_address,
            'arabic_work_address': arabic_work_address,  # New key added for Arabic work address
            'company_registrar': company_registrar,
            'company_country': company_country,
            'company_arabic_name': company_arabic_name,
            'head_people_culture': head_people_culture,
            'head_people_culture_arabic': head_people_culture_arabic
        }
    except Exception as e:
        st.error(f"Error retrieving employee data: {e}")
        return None

def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            replaced = True
    if not replaced:
        full_text = "".join(run.text for run in paragraph.runs)
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, replacement)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text

def remove_empty_paragraphs(doc: Document) -> None:
    for para in list(doc.paragraphs):
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)

def fill_template(template_path: str, employee_data: Dict[str, Any], is_arabic: bool = False) -> Optional[bytes]:
    if not os.path.exists(template_path):
        st.error(f"Template file not found: {template_path}")
        return None
    try:
        doc = Document(template_path)
    except Exception as e:
        st.error(f"Error loading document: {e}")
        return None

    current_date = datetime.date.today().strftime("%d/%m/%Y")
    placeholders = {
        "(Current Date)": current_date,
        "(First and Last Name)": employee_data['name'],
        "(First Name)": employee_data['first_name'],
        "(Position)": employee_data['job_title'],
        "(Salary)": str(employee_data['wage']),
        "(DD/MM/YYYY)": employee_data['joining_date'],
        "(Country)": employee_data.get('country', ''),
        "(Start Date)": employee_data.get('start_date', ''),
        "(End Date)": employee_data.get('end_date', ''),
        "(Company)": employee_data.get('company', ''),
        "(Work address)": employee_data.get('work_address', ''),
        "(Work Address)": employee_data.get('work_address', ''),
        "(Arabic Work address)": employee_data.get('arabic_work_address', ''),  # New placeholder added
        "(CR)": employee_data.get('company_registrar', ''),
        "(Company Country)": employee_data.get('company_country', ''),
        "(CompanyA)": employee_data.get('company_arabic_name', ''),
        "(P&C)": employee_data.get('head_people_culture', ''),
        "(AP&C)": employee_data.get('head_people_culture_arabic', ''),
        "(الاسم الكامل)": employee_data.get("arabic_name", employee_data['name']) if is_arabic else employee_data['name'],
        "(بلد الوجهة)": employee_data.get('country', ''),
        "(تاريخ البداية)": employee_data.get('start_date', ''),
        "(تاريخ النهاية)": employee_data.get('end_date', ''),
        "(Contract End Date)": employee_data.get('contract_end_date', ''),
        "(Department)": employee_data.get('department', '')
    }
    placeholders = {k: str(v) for k, v in placeholders.items()}
    
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_placeholder_in_paragraph(para, key, value)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        replace_placeholder_in_paragraph(para, key, value)
    
    for section in doc.sections:
        for para in section.header.paragraphs:
            for key, value in placeholders.items():
                replace_placeholder_in_paragraph(para, key, value)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in placeholders.items():
                            replace_placeholder_in_paragraph(para, key, value)
    
    for section in doc.sections:
        for para in section.footer.paragraphs:
            text = para.text
            for key, value in placeholders.items():
                if key in text:
                    text = text.replace(key, value)
            para.text = text
            for run in para.runs:
                run.font.size = Pt(8)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        for key, value in placeholders.items():
                            if key in text:
                                text = text.replace(key, value)
                        para.text = text
                        for run in para.runs:
                            run.font.size = Pt(8)
    
    for para in doc.paragraphs:
        if "(CR)" in para.text:
            para.text = para.text.replace("(CR)", employee_data.get("company_registrar", ""))
        if "(Company Country)" in para.text:
            para.text = para.text.replace("(Company Country)", employee_data.get("company_country", ""))
        if "(CompanyA)" in para.text:
            para.text = para.text.replace("(CompanyA)", employee_data.get("company_arabic_name", ""))
        if "(P&C)" in para.text:
            para.text = para.text.replace("(P&C)", employee_data.get("head_people_culture", ""))
        if "(AP&C)" in para.text:
            para.text = para.text.replace("(AP&C)", employee_data.get("head_people_culture_arabic", ""))
    
    remove_empty_paragraphs(doc)
    
    output_stream = io.BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

# List of countries (a fairly comprehensive list)
countries = [
    "Afghanistan", "Albania", "Algeria", "Andorra", "Angola", "Antigua and Barbuda",
    "Argentina", "Armenia", "Australia", "Austria", "Azerbaijan", "Bahamas", "Bahrain",
    "Bangladesh", "Barbados", "Belarus", "Belgium", "Belize", "Benin", "Bhutan", "Bolivia",
    "Bosnia and Herzegovina", "Botswana", "Brazil", "Brunei", "Bulgaria", "Burkina Faso",
    "Burundi", "Cabo Verde", "Cambodia", "Cameroon", "Canada", "Central African Republic",
    "Chad", "Chile", "China", "Colombia", "Comoros", "Congo (Congo-Brazzaville)", "Costa Rica",
    "Croatia", "Cuba", "Cyprus", "Czechia (Czech Republic)", "Democratic Republic of the Congo",
    "Denmark", "Djibouti", "Dominica", "Dominican Republic", "Ecuador", "Egypt", "El Salvador",
    "Equatorial Guinea", "Eritrea", "Estonia", "Eswatini (fmr. \"Swaziland\")", "Ethiopia",
    "Fiji", "Finland", "France", "Gabon", "Gambia", "Georgia", "Germany", "Ghana", "Greece",
    "Grenada", "Guatemala", "Guinea", "Guinea-Bissau", "Guyana", "Haiti", "Holy See", "Honduras",
    "Hungary", "Iceland", "India", "Indonesia", "Iran", "Iraq", "Ireland", "Italy",
    "Jamaica", "Japan", "Jordan", "Kazakhstan", "Kenya", "Kiribati", "Kuwait", "Kyrgyzstan",
    "Laos", "Latvia", "Lebanon", "Lesotho", "Liberia", "Libya", "Liechtenstein", "Lithuania",
    "Luxembourg", "Madagascar", "Malawi", "Malaysia", "Maldives", "Mali", "Malta",
    "Marshall Islands", "Mauritania", "Mauritius", "Mexico", "Micronesia", "Moldova", "Monaco",
    "Mongolia", "Montenegro", "Morocco", "Mozambique", "Myanmar (formerly Burma)", "Namibia",
    "Nauru", "Nepal", "Netherlands", "New Zealand", "Nicaragua", "Niger", "Nigeria", "North Korea",
    "North Macedonia", "Norway", "Oman", "Pakistan", "Palau", "Palestine State", "Panama",
    "Papua New Guinea", "Paraguay", "Peru", "Philippines", "Poland", "Portugal", "Qatar",
    "Romania", "Russia", "Rwanda", "Saint Kitts and Nevis", "Saint Lucia",
    "Saint Vincent and the Grenadines", "Samoa", "San Marino", "Sao Tome and Principe",
    "Saudi Arabia", "Senegal", "Serbia", "Seychelles", "Sierra Leone", "Singapore", "Slovakia",
    "Slovenia", "Solomon Islands", "Somalia", "South Africa", "South Korea", "South Sudan",
    "Spain", "Sri Lanka", "Sudan", "Suriname", "Sweden", "Switzerland", "Syria", "Tajikistan",
    "Tanzania", "Thailand", "Timor-Leste", "Togo", "Tonga", "Trinidad and Tobago", "Tunisia",
    "Turkey", "Turkmenistan", "Tuvalu", "Uganda", "Ukraine", "United Arab Emirates", "United Kingdom",
    "United States of America", "Uruguay", "Uzbekistan", "Vanuatu", "Venezuela", "Vietnam", "Yemen",
    "Zambia", "Zimbabwe"
]

st.set_page_config(
    page_title="Employment Letter Generator",
    page_icon=":briefcase:",
    layout="centered"
)

st.markdown(
    """
    <style>
    .main { padding: 2rem; }
    .stButton>button { background-color: #2e7bcf; color: white; border-radius: 5px; }
    </style>
    """,
    unsafe_allow_html=True
)

# Template options (files are expected to be in the same directory as this code).
template_options = {
    "Employment letter - Arabic": "Employment Letter - ARABIC 1.docx",
    "Employment letter": "Employment Letter 1.docx",
    "Employment letter to embassies": "Employment Letter to Embassies 1.docx",
    "Experience letter": "Experience Letter 1.docx"
}

def main() -> None:
    st.title("Employment Letter Generator 🚀")
    st.markdown("Please fill in the details below to generate the employment letter. ✨")
    
    # Template selection outside the form.
    selected_template = st.selectbox("Select Template", list(template_options.keys()))
    template_path = template_options[selected_template]
    
    if "employee_id" not in st.session_state:
        st.session_state["employee_id"] = ""
    
    with st.form("letter_form"):
        identification_no = st.text_input("Employee Identification Number 🆔", value=st.session_state["employee_id"])
        
        # For embassy letters, show embassy travel details; country is now a dropdown.
        country = ""
        start_date = None
        end_date = None
        if selected_template == "Employment letter to embassies":
            with st.expander("Embassy Travel Details", expanded=True):
                col1, col2, col3 = st.columns(3)
                country = col1.selectbox("Country Name 🌍", options=countries)
                start_date = col2.date_input("Travel Start Date 📆")
                end_date = col3.date_input("Travel End Date 📆")
        
        submitted = st.form_submit_button("Generate Letter ✨")
        if submitted:
            st.session_state["employee_id"] = identification_no
            uid, models = get_odoo_connection()
            if not uid:
                st.error("Could not connect to Odoo.")
                st.stop()
            employee_data = get_employee_by_id(models, uid, identification_no)
            if not employee_data:
                st.error("Could not retrieve employee data.")
                st.stop()
            if selected_template == "Employment letter to embassies":
                employee_data['country'] = country.strip()
                employee_data['start_date'] = start_date.strftime("%d/%m/%Y") if start_date else ""
                employee_data['end_date'] = end_date.strftime("%d/%m/%Y") if end_date else ""
            else:
                employee_data['country'] = ""
                employee_data['start_date'] = ""
                employee_data['end_date'] = ""
            
            st.markdown("### Employee Details")
            st.write(f"**Name:** {employee_data.get('name', '')}")
            st.write(f"**Job Title:** {employee_data.get('job_title', '')}")
            st.write(f"**Joining Date:** {employee_data.get('joining_date', '')}")
            st.write(f"**Contract End Date:** {employee_data.get('contract_end_date', '')}")
            st.write(f"**Department:** {employee_data.get('department', '')}")
            st.write(f"**Wage:** {employee_data.get('wage', '')}")
            st.write(f"**Company:** {employee_data.get('company', '')}")
            st.write(f"**Work Address:** {employee_data.get('work_address', '')}")
            st.write(f"**Arabic Work Address:** {employee_data.get('arabic_work_address', '')}")
            st.write(f"**Company Registrar (CR):** {employee_data.get('company_registrar', '')}")
            st.write(f"**Company Country:** {employee_data.get('company_country', '')}")
            st.write(f"**Company Arabic Name (CompanyA):** {employee_data.get('company_arabic_name', '')}")
            st.write(f"**Head of People & Culture (P&C):** {employee_data.get('head_people_culture', '')}")
            st.write(f"**Head of People & Culture Arabic (AP&C):** {employee_data.get('head_people_culture_arabic', '')}")
            
            # File naming: [Template Name] - [Employee Name].docx
            filename = f"{selected_template} - {employee_data['name']}.docx"
            doc_bytes = fill_template(template_path, employee_data, is_arabic=(selected_template == "Employment letter - Arabic"))
            if not doc_bytes:
                st.error("Failed to generate the document.")
            else:
                st.session_state["letter_bytes"] = doc_bytes
                st.session_state["letter_filename"] = filename
                st.success("Employment Letter Generated Successfully! 🎉")
    
    if st.session_state.get("letter_bytes"):
        st.download_button(
            label="Download Employment Letter 📄",
            data=st.session_state["letter_bytes"],
            file_name=st.session_state["letter_filename"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
